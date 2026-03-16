"""
diffcheck.py — Compare two Word documents and produce a diff .docx.

Removed words appear in RED with STRIKETHROUGH.
Added words appear in GREEN.
Unchanged words appear in the default style.

Usage (called by Next.js API route):
    python3 diffcheck.py <original.docx> <revised.docx> <output.docx>

Usage (legacy / standalone):
    python3 diffcheck.py
    # compares SampleDoc.docx vs SampleDoc1.docx → SampleDoc_diff.docx
"""

import sys
import difflib
from docx import Document
from docx.shared import RGBColor


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def extract_paragraphs(doc_path: str) -> list:
    """Return a list of paragraphs; each paragraph is a list of word tokens."""
    doc = Document(doc_path)
    result = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            result.append(text.split())
    return result


def diff_word_lists(old_words: list, new_words: list):
    """
    Run a SequenceMatcher diff on two word lists.
    Yields tuples of (tag, word) where tag is one of:
        'equal'   - unchanged
        'delete'  - in old only  (render red + strikethrough)
        'insert'  - in new only  (render green)
    """
    matcher = difflib.SequenceMatcher(None, old_words, new_words, autojunk=False)
    for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
        if opcode == "equal":
            for w in old_words[i1:i2]:
                yield ("equal", w)
        elif opcode == "delete":
            for w in old_words[i1:i2]:
                yield ("delete", w)
        elif opcode == "insert":
            for w in new_words[j1:j2]:
                yield ("insert", w)
        elif opcode == "replace":
            for w in old_words[i1:i2]:
                yield ("delete", w)
            for w in new_words[j1:j2]:
                yield ("insert", w)


# ---------------------------------------------------------------------------
# Core
# ---------------------------------------------------------------------------

def generate_diff_doc(original_path: str, revised_path: str, output_path: str):
    orig_paras = extract_paragraphs(original_path)
    rev_paras  = extract_paragraphs(revised_path)

    # Align paragraphs by index; pad the shorter list with empty paragraphs
    max_len = max(len(orig_paras), len(rev_paras))
    orig_paras += [[]] * (max_len - len(orig_paras))
    rev_paras  += [[]] * (max_len - len(rev_paras))

    diff_doc = Document()

    stats = {"equal": 0, "delete": 0, "insert": 0}

    for old_words, new_words in zip(orig_paras, rev_paras):
        paragraph = diff_doc.add_paragraph()

        for tag, word in diff_word_lists(old_words, new_words):
            stats[tag] += 1
            run = paragraph.add_run(word + " ")

            if tag == "delete":
                run.font.color.rgb = RGBColor(192, 57, 43)   # red
                run.font.strike    = True
            elif tag == "insert":
                run.font.color.rgb = RGBColor(39, 174, 96)   # green
            # equal -> default formatting

    diff_doc.save(output_path)

    print(
        "[diffcheck] Done -- "
        + str(stats["equal"]) + " unchanged, "
        + str(stats["delete"]) + " removed, "
        + str(stats["insert"]) + " added -> " + output_path
    )

    # Return stats so the API route can surface them
    return stats


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    args = sys.argv[1:]

    if len(args) == 3:
        original_path, revised_path, output_path = args
    elif len(args) == 0:
        original_path = "SampleDoc.docx"
        revised_path  = "SampleDoc1.docx"
        output_path   = "SampleDoc_diff.docx"
    else:
        print(
            "Usage: python3 diffcheck.py [original.docx revised.docx output.docx]",
            file=sys.stderr,
        )
        sys.exit(1)

    generate_diff_doc(original_path, revised_path, output_path)
